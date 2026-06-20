import os
import io
from fastapi import FastAPI, HTTPException, Depends, status
from fastapi.responses import HTMLResponse, StreamingResponse
from fastapi.security import HTTPBasic, HTTPBasicCredentials
from pydantic import BaseModel
from google import genai
from dotenv import load_dotenv

# 📁 ファイル作成用のライブラリ
from docx import Document
import openpyxl

load_dotenv()

app = FastAPI()

# 🔐 パスワード認証の仕組み
security = HTTPBasic()

def authenticate(credentials: HTTPBasicCredentials = Depends(security)):
    correct_username = os.getenv("WEB_USERNAME", "admin")
    correct_password = os.getenv("WEB_PASSWORD", "password")
    if credentials.username != correct_username or credentials.password != correct_password:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="ユーザー名またはパスワードが違います",
            headers={"WWW-Authenticate": "Basic"},
        )
    return credentials.username

# Geminiクライアントの初期化
try:
    client = genai.Client()
except Exception as e:
    print(f"Geminiの初期化に失敗しました。APIキーを確認してください: {e}")

class TaskRequest(BaseModel):
    title: str
    description: str

# ダウンロード時にフロントから送られてくるデータの型
class DownloadRequest(BaseModel):
    title: str
    result: str

@app.get("/", response_class=HTMLResponse)
async def read_index(username: str = Depends(authenticate)):
    with open("index.html", "r", encoding="utf-8") as f:
        return f.read()

@app.post("/process-task")
async def process_task(request: TaskRequest, username: str = Depends(authenticate)):
# 📝 Geminiへの指示文をパワーアップ（言い訳禁止令を追加）
    prompt = f"""
    以下のタスクを分析し、効率よく終わらせるための「具体的な実行ステップ（やることリスト）」と「ワンポイントアドバイス」を丁寧な箇条書きで出力してください。
    
    【超重要ルール】
    ・出力されたテキストは、この後システムによって自動的にWord（.docx）やExcel（.xlsx）ファイルに変換されます。
    ・そのため、「私は直接ファイルを生成できません」「Wordに貼り付けてご利用ください」といった、AIとしての限界や言い訳、案内文は【絶対に】出力に含めないでください。
    ・前置きの挨拶や結びの言葉も一切不要です。純粋なタスクの分析結果（タイトル、ステップ、アドバイス）のコンテンツだけを出力してください。
    
    【タスク名】: {request.title}
    【詳細・メモ】: {request.description}
    """
    try:
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=prompt,
        )
        return {"status": "success", "result": response.text}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# 📄 ① テキストファイル（.txt）の生成・ダウンロード
@app.post("/download/txt")
async def download_txt(request: DownloadRequest, username: str = Depends(authenticate)):
    content = f"【タスク名】: {request.title}\n\n{request.result}"
    file_io = io.BytesIO(content.encode("utf-8"))
    return StreamingResponse(
        file_io,
        media_type="text/plain",
        headers={"Content-Disposition": "attachment; filename=task.txt"}
    )

# 📝 ② Wordファイル（.docx）の生成・ダウンロード
@app.post("/download/docx")
async def download_docx(request: DownloadRequest, username: str = Depends(authenticate)):
    doc = Document()
    doc.add_heading(f"タスク分析: {request.title}", level=1)
    
    # AIの返答を1行ずつ解析して綺麗にWordに成形
    for line in request.result.split("\n"):
        if line.startswith("### "):
            doc.add_heading(line.replace("### ", ""), level=3)
        elif line.startswith("## "):
            doc.add_heading(line.replace("## ", ""), level=2)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style='List Bullet')
        else:
            if line.strip():
                doc.add_paragraph(line)
                
    file_io = io.BytesIO()
    doc.save(file_io)
    file_io.seek(0)
    return StreamingResponse(
        file_io,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=task.docx"}
    )

# 📊 ③ Excelファイル（.xlsx）の生成・ダウンロード
@app.post("/download/xlsx")
async def download_xlsx(request: DownloadRequest, username: str = Depends(authenticate)):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "タスク詳細"
    
    ws.append(["タスク名", request.title])
    ws.append([])
    ws.append(["AI分析結果（ログ）"])
    
    for line in request.result.split("\n"):
        ws.append([line])
        
    file_io = io.BytesIO()
    wb.save(file_io)
    file_io.seek(0)
    return StreamingResponse(
        file_io,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=task.xlsx"}
    )